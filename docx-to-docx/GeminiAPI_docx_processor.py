import os
import json
from docx import Document
import requests

def load_api_config():
    """Load Gemini API configuration from a JSON file."""
    with open('docx-to-docx/gemini_api_config.json') as config_file:
        return json.load(config_file)

def read_docx(file_path):
    """Read and return the text from a DOCX file."""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def write_docx(content, output_path):
    """Write the given content to a DOCX file."""
    doc = Document()
    doc.add_paragraph(content)
    doc.save(output_path)

def process_docx_files(directory):
    """Process all DOCX files in the given directory with the Gemini API."""
    api_config = load_api_config()
    api_url = api_config['api_url']
    headers = {'Authorization': f"Bearer {api_config['api_key']}"}

    for file_name in os.listdir(directory):
        if file_name.endswith('.docx'):
            file_path = os.path.join(directory, file_name)
            text_content = read_docx(file_path)
            response = requests.post(api_url, headers=headers, json={'text': text_content})
            if response.status_code == 200:
                output_content = response.json()['response']
                output_file_path = os.path.join(directory, f"processed_{file_name}")
                write_docx(output_content, output_file_path)
                print(f"Processed and saved: {output_file_path}")
            else:
                print(f"Error processing file {file_name}: {response.text}")

if __name__ == "__main__":
    directory = input("Enter the directory path containing DOCX files to process: ")
    process_docx_files(directory)
